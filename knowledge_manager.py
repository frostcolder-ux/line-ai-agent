"""
Knowledge Base Manager
雙模式：有 DATABASE_URL → PostgreSQL（資料永久保存）；無 → 本機 JSON 檔。

向量語意搜尋（RAG）：
  - 需要設定環境變數 VOYAGE_API_KEY（免費申請：https://dash.voyageai.com）
  - 上傳文件時自動切分段落並建立 512 維嵌入向量
  - 查詢時只取最相關的 TOP_K 段落，不受知識庫大小限制
  - 未設定 API Key 時自動降級為截斷全文模式
"""
import os
import json
import uuid
import sys
from datetime import datetime

BASE_DIR = os.path.dirname(__file__)
KB_DIR   = os.path.join(BASE_DIR, "knowledge_db")
KB_PATH  = os.path.join(KB_DIR, "db.json")

_CHUNK_SIZE    = 500   # 每段最大字元數
_CHUNK_OVERLAP = 80    # 相鄰段落重疊字元數，保留語境
_TOP_K         = 6     # 每次查詢取最相關幾段
_KB_MAX_CHARS  = 80_000  # 無向量搜尋時的截斷上限


def log(msg):
    print(f"[KB] {msg}", file=sys.stderr, flush=True)


# ══════════════════════════════════════════════════════════════════════════════
# 向量嵌入（Voyage AI）
# ══════════════════════════════════════════════════════════════════════════════

def _get_voyage_client():
    """取得 Voyage AI 客戶端；無 API Key 時回傳 None。"""
    api_key = os.environ.get("VOYAGE_API_KEY", "").strip()
    if not api_key:
        return None
    try:
        import voyageai
        return voyageai.Client(api_key=api_key)
    except ImportError:
        log("voyageai 套件未安裝")
        return None


def has_vector_search() -> bool:
    """判斷向量搜尋是否可用（需要 VOYAGE_API_KEY + PostgreSQL）。"""
    from db import is_postgres
    return bool(os.environ.get("VOYAGE_API_KEY", "").strip()) and is_postgres()


def _embed(texts: list[str], input_type: str = "document") -> list[list[float]] | None:
    """
    呼叫 Voyage AI 產生嵌入向量。
    input_type: "document"（上傳時）或 "query"（搜尋時）
    """
    client = _get_voyage_client()
    if not client or not texts:
        return None
    try:
        result = client.embed(texts, model="voyage-3-lite", input_type=input_type)
        return result.embeddings
    except Exception as e:
        log(f"Voyage AI 嵌入失敗: {e}")
        return None


def _vec_str(embedding: list[float]) -> str:
    """把 Python list 轉成 pgvector 接受的字串格式：'[0.1,0.2,...]'。"""
    return "[" + ",".join(f"{x:.6f}" for x in embedding) + "]"


# ══════════════════════════════════════════════════════════════════════════════
# 文件切段
# ══════════════════════════════════════════════════════════════════════════════

def _extract_keywords(text: str) -> list[str]:
    """
    從問題中提取關鍵詞：
    - 中文：取 2 字以上的詞語 + 2-gram 切分
    - 英文：取 3 字母以上的單詞
    - 過濾常見虛詞/停用字
    """
    import re
    STOPWORDS = {
        '什麼', '如何', '為何', '為什麼', '怎麼', '怎樣', '可以', '是否',
        '請問', '告訴', '我要', '想知', '知道', '幫我', '這個', '那個',
        '有沒', '有沒有', '一些', '應該', '需要', '可能', '是的', '對嗎',
        '嗎', '呢', '啊', '吧', '喔', '哦', '嗯',
    }
    clean   = re.sub(r'[^\w一-鿿]', ' ', text)
    tokens  = []
    # 中文詞語
    for word in re.findall(r'[一-鿿]+', clean):
        if len(word) >= 2 and word not in STOPWORDS:
            tokens.append(word)
            for i in range(len(word) - 1):           # 2-gram
                gram = word[i:i+2]
                if gram not in STOPWORDS:
                    tokens.append(gram)
    # 英文單詞
    for word in re.findall(r'[a-zA-Z]{3,}', clean):
        tokens.append(word.lower())
    return list(set(tokens))


def chunk_text(text: str) -> list[str]:
    """
    把長文字切成有重疊的小段落，適合關鍵字搜尋與嵌入。
    優先在段落/句子邊界切斷，保持語義完整。
    """
    chunks = []
    start  = 0
    length = len(text)
    while start < length:
        end   = min(start + _CHUNK_SIZE, length)
        chunk = text[start:end]
        # 嘗試在自然邊界切斷
        if end < length:
            for sep in ["\n\n", "\n", "。", "！", "？", ".", "!", "?", "，", ","]:
                pos = chunk.rfind(sep)
                if pos > _CHUNK_SIZE // 3:
                    chunk = chunk[:pos + len(sep)]
                    break
        chunk = chunk.strip()
        if chunk:
            chunks.append(chunk)
        advance = len(chunk) - _CHUNK_OVERLAP
        start  += max(advance, 1)
    return chunks


# ══════════════════════════════════════════════════════════════════════════════
# PostgreSQL 向量 chunks 操作
# ══════════════════════════════════════════════════════════════════════════════

def _pg_store_chunks(doc_id: str, filename: str, content: str):
    """
    切段文件並儲存到 knowledge_chunks。
    - 有 VOYAGE_API_KEY → 同時儲存嵌入向量（向量搜尋）
    - 無 VOYAGE_API_KEY → 只儲存文字（關鍵字搜尋仍可用）
    """
    chunks = chunk_text(content)
    if not chunks:
        return 0

    # 嘗試建立向量（失敗不影響儲存）
    embeddings = _embed(chunks, input_type="document")  # None 表示無 API Key

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    from db import get_conn
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute("DELETE FROM knowledge_chunks WHERE doc_id = %s", (doc_id,))
    for i, chunk in enumerate(chunks):
        emb = embeddings[i] if embeddings else None
        if emb is not None:
            cur.execute(
                "INSERT INTO knowledge_chunks (doc_id, filename, chunk_idx, content, embedding, created_at) "
                "VALUES (%s, %s, %s, %s, %s::vector, %s)",
                (doc_id, filename, i, chunk, _vec_str(emb), now)
            )
        else:
            cur.execute(
                "INSERT INTO knowledge_chunks (doc_id, filename, chunk_idx, content, created_at) "
                "VALUES (%s, %s, %s, %s, %s)",
                (doc_id, filename, i, chunk, now)
            )
    conn.commit()
    cur.close()
    conn.close()
    mode = "含向量" if embeddings else "純文字（關鍵字搜尋）"
    log(f"儲存 {len(chunks)} 段（{mode}）到 DB（{filename}）")
    return len(chunks)


def _pg_delete_chunks(doc_id: str):
    from db import get_conn
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute("DELETE FROM knowledge_chunks WHERE doc_id = %s", (doc_id,))
    conn.commit()
    cur.close()
    conn.close()


def _pg_clear_chunks():
    from db import get_conn
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute("DELETE FROM knowledge_chunks")
    conn.commit()
    cur.close()
    conn.close()


def _pg_search_chunks(query: str, top_k: int = _TOP_K) -> list[dict]:
    """用餘弦相似度找出最相關的段落，回傳 [{filename, content, similarity}]。"""
    embeddings = _embed([query], input_type="query")
    if not embeddings:
        return []
    vec = _vec_str(embeddings[0])
    from db import get_conn
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute(
        """
        SELECT filename, content, 1 - (embedding <=> %s::vector) AS similarity
        FROM knowledge_chunks
        WHERE embedding IS NOT NULL
        ORDER BY embedding <=> %s::vector
        LIMIT %s
        """,
        (vec, vec, top_k)
    )
    rows = cur.fetchall()
    cur.close()
    conn.close()
    return [{"filename": r[0], "content": r[1], "similarity": r[2]} for r in rows]


def _pg_keyword_search(query: str, top_k: int = _TOP_K) -> list[dict]:
    """
    關鍵字搜尋：從 knowledge_chunks 找出包含最多關鍵字的段落。
    不需要 Voyage AI，純 SQL，完全免費。
    """
    keywords = _extract_keywords(query)
    if not keywords:
        return []

    from db import get_conn
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute("SELECT filename, content FROM knowledge_chunks ORDER BY doc_id, chunk_idx")
    rows = cur.fetchall()
    cur.close()
    conn.close()

    scored = []
    for filename, content in rows:
        # 計算命中幾個關鍵字（每個關鍵字最多算一次）
        score = sum(1 for kw in keywords if kw in content)
        if score > 0:
            scored.append((score, filename, content))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [{"filename": r[1], "content": r[2], "score": r[0]} for r in scored[:top_k]]


def _json_keyword_search(query: str, top_k: int = _TOP_K) -> str:
    """
    JSON 模式的關鍵字搜尋：在記憶體中即時切段並比對（無 DB 時使用）。
    """
    keywords = _extract_keywords(query)
    if not keywords:
        content = _json_get_all_content()
        return content[:_KB_MAX_CHARS] if len(content) > _KB_MAX_CHARS else content

    db     = _json_load_kb()
    scored = []
    for doc in db.get("documents", []):
        for chunk in chunk_text(doc["content"]):
            score = sum(1 for kw in keywords if kw in chunk)
            if score > 0:
                scored.append((score, doc["filename"], chunk))

    scored.sort(key=lambda x: x[0], reverse=True)
    top = scored[:top_k]
    if not top:
        return ""
    parts = [f"### 來源：{r[1]}（命中 {r[0]} 個關鍵字）\n{r[2]}" for r in top]
    return "\n\n---\n\n".join(parts)


def has_keyword_search() -> bool:
    """關鍵字搜尋只需要有 chunks（PostgreSQL 模式自動儲存）。"""
    from db import is_postgres
    if not is_postgres():
        return True   # JSON 模式也支援，直接在記憶體切段搜尋
    try:
        from db import get_conn
        conn = get_conn()
        cur  = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM knowledge_chunks")
        count = cur.fetchone()[0]
        cur.close()
        conn.close()
        return count > 0
    except Exception:
        return False


def reembed_all() -> dict:
    """重新對所有文件建立 chunks（+ 向量，若有 API Key）。"""
    from db import is_postgres
    if not is_postgres():
        return {"error": "僅支援 PostgreSQL 模式"}

    kb   = _pg_load_kb()
    docs = kb.get("documents", [])
    total_chunks = 0
    for doc in docs:
        n = _pg_store_chunks(doc["id"], doc["filename"], doc["content"])
        total_chunks += n
    mode = "向量+關鍵字" if has_vector_search() else "關鍵字"
    return {"docs": len(docs), "chunks": total_chunks, "mode": mode}


# ══════════════════════════════════════════════════════════════════════════════
# PostgreSQL 實作
# ══════════════════════════════════════════════════════════════════════════════

def _pg_load_kb() -> dict:
    from db import get_conn
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute(
        "SELECT id, filename, content, chars, uploaded_at, preview "
        "FROM knowledge_docs ORDER BY uploaded_at"
    )
    rows = cur.fetchall()
    cur.close(); conn.close()
    docs = [
        {"id": r[0], "filename": r[1], "content": r[2],
         "chars": r[3], "uploaded_at": r[4], "preview": r[5]}
        for r in rows
    ]
    return {"documents": docs}


def _pg_add_document(filename: str, content: str) -> dict:
    from db import get_conn
    doc_id  = str(uuid.uuid4())[:8]
    now     = datetime.now().strftime("%Y-%m-%d %H:%M")
    preview = content[:200].strip()
    chars   = len(content)
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute(
        "INSERT INTO knowledge_docs (id, filename, content, chars, uploaded_at, preview) "
        "VALUES (%s, %s, %s, %s, %s, %s)",
        (doc_id, filename, content, chars, now, preview)
    )
    conn.commit(); cur.close(); conn.close()
    log(f"PG: added {filename} ({chars} chars)")
    # 建立向量索引（非同步不影響上傳速度，失敗不影響主流程）
    try:
        _pg_store_chunks(doc_id, filename, content)
    except Exception as e:
        log(f"PG: chunk 建立失敗（不影響上傳）: {e}")
    return {"id": doc_id, "filename": filename, "content": content,
            "chars": chars, "uploaded_at": now, "preview": preview}


def _pg_delete_document(doc_id: str) -> bool:
    from db import get_conn
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute("DELETE FROM knowledge_docs WHERE id = %s", (doc_id,))
    deleted = cur.rowcount > 0
    conn.commit(); cur.close(); conn.close()
    if deleted:
        try:
            _pg_delete_chunks(doc_id)
        except Exception as e:
            log(f"PG: 刪除 chunks 失敗: {e}")
    return deleted


def _pg_clear_all() -> int:
    from db import get_conn
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM knowledge_docs")
    count = cur.fetchone()[0]
    cur.execute("DELETE FROM knowledge_docs")
    conn.commit(); cur.close(); conn.close()
    try:
        _pg_clear_chunks()
    except Exception as e:
        log(f"PG: 清除 chunks 失敗: {e}")
    return count


def _pg_get_all_content() -> str:
    from db import get_conn
    conn = get_conn()
    cur  = conn.cursor()
    cur.execute("SELECT filename, content FROM knowledge_docs ORDER BY uploaded_at")
    rows = cur.fetchall()
    cur.close(); conn.close()
    if not rows:
        return ""
    parts = [f"### 文件來源：{r[0]}\n{r[1]}" for r in rows]
    return "\n\n---\n\n".join(parts)


# ══════════════════════════════════════════════════════════════════════════════
# JSON 檔實作（本地開發 / 無 DB 備援）
# ══════════════════════════════════════════════════════════════════════════════

def _json_load_kb() -> dict:
    try:
        with open(KB_PATH, encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {"documents": []}


def _json_save_kb(db: dict):
    os.makedirs(KB_DIR, exist_ok=True)
    with open(KB_PATH, "w", encoding="utf-8") as f:
        json.dump(db, f, ensure_ascii=False, indent=2)


def _json_add_document(filename: str, content: str) -> dict:
    db  = _json_load_kb()
    doc = {
        "id":          str(uuid.uuid4())[:8],
        "filename":    filename,
        "content":     content,
        "chars":       len(content),
        "uploaded_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "preview":     content[:200].strip(),
    }
    db["documents"].append(doc)
    _json_save_kb(db)
    return doc


def _json_delete_document(doc_id: str) -> bool:
    db     = _json_load_kb()
    before = len(db["documents"])
    db["documents"] = [d for d in db["documents"] if d["id"] != doc_id]
    if len(db["documents"]) < before:
        _json_save_kb(db)
        return True
    return False


def _json_clear_all() -> int:
    db    = _json_load_kb()
    count = len(db["documents"])
    _json_save_kb({"documents": []})
    return count


def _json_get_all_content() -> str:
    db = _json_load_kb()
    if not db["documents"]:
        return ""
    parts = [f"### 文件來源：{d['filename']}\n{d['content']}" for d in db["documents"]]
    return "\n\n---\n\n".join(parts)


# ══════════════════════════════════════════════════════════════════════════════
# 公開 API（自動選擇模式）
# ══════════════════════════════════════════════════════════════════════════════

def _use_pg() -> bool:
    from db import is_postgres
    return is_postgres()


def load_kb() -> dict:
    return _pg_load_kb() if _use_pg() else _json_load_kb()


def save_kb(db: dict):
    """匯入備份用（只支援 JSON 模式，PG 模式請用 add_document）"""
    if _use_pg():
        _pg_clear_all()
        for doc in db.get("documents", []):
            _pg_add_document(doc["filename"], doc["content"])
    else:
        _json_save_kb(db)


def add_document(filename: str, content: str) -> dict:
    return _pg_add_document(filename, content) if _use_pg() else _json_add_document(filename, content)


def delete_document(doc_id: str) -> bool:
    return _pg_delete_document(doc_id) if _use_pg() else _json_delete_document(doc_id)


def clear_all() -> int:
    return _pg_clear_all() if _use_pg() else _json_clear_all()


def get_all_content() -> str:
    try:
        return _pg_get_all_content() if _use_pg() else _json_get_all_content()
    except Exception as e:
        log(f"get_all_content error: {e}")
        return ""


def search_relevant_chunks(query: str, top_k: int = _TOP_K) -> str:
    """
    根據問題找出最相關的知識庫段落，回傳格式化文字供 system prompt 使用。

    搜尋策略（依優先順序）：
    1. 向量語意搜尋（需要 VOYAGE_API_KEY）  ← 最準確
    2. 關鍵字搜尋（免費，自動可用）          ← 中等準確，無上限
    3. 截斷全文（最後手段，限 80,000 字元）   ← 保底

    搜不到任何相關資料時回傳空字串 → build_system_prompt 不注入知識庫
    → Claude 會說「查無相關資料」而不是亂回答。
    """
    if not query.strip():
        content = get_all_content()
        return content[:_KB_MAX_CHARS] if len(content) > _KB_MAX_CHARS else content

    # ── 第一層：向量語意搜尋 ────────────────────────────────────────────────
    if has_vector_search():
        try:
            results = _pg_search_chunks(query, top_k=top_k)
            if results:
                parts = [
                    f"### 來源：{r['filename']}（相關度 {r['similarity']:.0%}）\n{r['content']}"
                    for r in results
                ]
                log(f"[搜尋] 向量：{len(results)} 段，最高 {results[0]['similarity']:.0%}")
                return "\n\n---\n\n".join(parts)
            log("[搜尋] 向量：無結果，降級關鍵字")
        except Exception as e:
            log(f"[搜尋] 向量失敗：{e}，降級關鍵字")

    # ── 第二層：關鍵字搜尋（免費，PG chunks 或 JSON 記憶體）─────────────────
    if _use_pg():
        try:
            results = _pg_keyword_search(query, top_k=top_k)
            if results:
                parts = [
                    f"### 來源：{r['filename']}（命中 {r['score']} 個關鍵字）\n{r['content']}"
                    for r in results
                ]
                log(f"[搜尋] 關鍵字：{len(results)} 段")
                return "\n\n---\n\n".join(parts)
            log("[搜尋] 關鍵字：無命中 → 回傳空（不注入知識庫）")
            return ""   # ← 無相關資料，讓 Claude 說「查無資料」
        except Exception as e:
            log(f"[搜尋] 關鍵字失敗：{e}，降級全文")
    else:
        result = _json_keyword_search(query, top_k=top_k)
        if result:
            log(f"[搜尋] JSON 關鍵字：找到相關段落")
            return result
        log("[搜尋] JSON 關鍵字：無命中 → 回傳空")
        return ""

    # ── 第三層：截斷全文（最後手段）────────────────────────────────────────
    content = get_all_content()
    if len(content) > _KB_MAX_CHARS:
        content = content[:_KB_MAX_CHARS]
        last_break = content.rfind("\n\n")
        if last_break > _KB_MAX_CHARS // 2:
            content = content[:last_break]
    log(f"[搜尋] 全文截斷：{len(content)} 字元")
    return content


def get_kb_stats() -> dict:
    db   = load_kb()
    docs = db["documents"]
    total_chars = sum(d.get("chars", 0) for d in docs)
    return {
        "document_count":   len(docs),
        "total_chars":      total_chars,
        "estimated_tokens": total_chars // 4,
        "documents":        docs,
    }


# ══════════════════════════════════════════════════════════════════════════════
# 文字解析
# ══════════════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {"txt", "md", "pdf", "docx", "csv"}


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支援的格式 .{ext}（支援：PDF、DOCX、TXT、MD、CSV）")

    if ext in ("txt", "md"):
        for enc in ("utf-8", "utf-8-sig", "big5", "gbk"):
            try:
                return file_bytes.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return file_bytes.decode("utf-8", errors="replace")

    if ext == "pdf":
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(file_bytes))
            pages  = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[第 {i+1} 頁]\n{text.strip()}")
            if not pages:
                raise ValueError("PDF 中未能提取到文字（可能是掃描圖檔，需 OCR）")
            return "\n\n".join(pages)
        except ImportError:
            raise ValueError("伺服器尚未安裝 pypdf")
        except Exception as e:
            raise ValueError(f"PDF 解析失敗：{e}")

    if ext == "docx":
        try:
            from docx import Document
            import io
            doc   = Document(io.BytesIO(file_bytes))
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            table_rows = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        table_rows.append(" | ".join(cells))
            all_text = paras + (["[表格資料]"] + table_rows if table_rows else [])
            return "\n\n".join(all_text)
        except ImportError:
            raise ValueError("伺服器尚未安裝 python-docx")
        except Exception as e:
            raise ValueError(f"DOCX 解析失敗：{e}")

    if ext == "csv":
        for enc in ("utf-8-sig", "utf-8", "big5"):
            try:
                return file_bytes.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return file_bytes.decode("utf-8", errors="replace")

    return ""
